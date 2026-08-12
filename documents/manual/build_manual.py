"""
Build the Van Sale user manual: van-sale-user-manual.md -> van-sale-user-manual.docx

Run with the 3.14 interpreter, which is the only one on this machine carrying
python-docx and Pillow:

    py -3.14 documents/manual/build_manual.py

Content lives in the Markdown file and *only* there. Every formatting decision
lives here. The numbers in SPEC below were measured out of the house reference
document by unzipping it and reading the WordprocessingML directly, so this is
the spec -- verify_manual.py imports SPEC from this module and asserts the
generated .docx against the same constants, which is what stops the spec and
the checker drifting apart.

Why python-docx and not an HTML -> Word conversion: run shading (the step
badge), negative paragraph indents (the part banner bleeding into the margin),
repeating table header rows, hanging-indent bullets without a numbering
definition, and multi-paragraph footers all survive this path and none of them
survive that one.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Emu, Inches, Pt, RGBColor, Twips
from PIL import Image

HERE = Path(__file__).resolve().parent
MD = HERE / "van-sale-user-manual.md"
DOCX = HERE / "van-sale-user-manual.docx"
SHOTS = HERE / "screenshots"
MANIFEST = SHOTS / "manifest.json"

PRODUCT = "Van Sale"
BRAND = "369ai.Biz"


# --------------------------------------------------------------------------
# SPEC - every measurement, in the units Word stores them in.
#
# Twips unless the name says otherwise (1440 twips = 1 inch = 20 half-points
# of font size). Colours are the raw hex Word writes into w:color / w:fill.
# --------------------------------------------------------------------------

SPEC = {
    # Page. US Letter, one inch of margin all round, half an inch of
    # header/footer reserve. Content width is 12240 - 1440 - 1440 = 9360, and
    # every full-width construct below sums to exactly that.
    "page_w": 12240,
    "page_h": 15840,
    "margin": 1440,
    "hf_distance": 720,
    "content_w": 9360,

    # Palette.
    "navy": "1F3864",        # title, table header text, closing banner text
    "banner_fill": "2E5395",  # part banner background and its hairline border
    "accent": "2E74B5",       # headings, step badge, brand
    "rule": "4A86C8",         # cover rules, header underline, footer overline
    "tbl_border": "B7CCE4",
    "tbl_head_fill": "EAF1F8",
    "tbl_col1": "1B3A5C",
    "body": "1F2937",
    "grey": "595959",
    "grey_dark": "404040",
    "grey_page": "262626",
    "fig_border": "AFC6E0",
    "fig_fill": "F7F9FC",
    "callout_border": "D97706",
    "callout_fill": "FFF8E7",

    # Type. Word stores sizes in half-points, so 22 = 11pt.
    "font": "Segoe UI",
    "sz_body": 22,
    "sz_small": 20,     # tables, captions, header, footer, step badge
    "sz_title": 72,     # cover title, 36pt
    "sz_subtitle": 40,  # cover subtitle, 20pt
    "sz_tagline": 24,
    "sz_brand": 28,
    "sz_banner": 32,    # part banner text, 16pt
    "sz_heading": 30,   # h2 and step title, 15pt
    "sz_closing": 21,

    # Constructs.
    "banner_indent": -86,      # bleeds the banner past the text edge
    "bullet_indent": 605,
    "bullet_hanging": 317,
    "bullet_glyph": "▪",
    "fig_cell_w": 6624,
    "fig_max_w_in": 3.2,       # images are scaled to fit inside this box
    "fig_max_h_in": 2.8,
    "callout_border_sz": 30,   # eighths of a point, so a 3.75pt left rule
    "grids": {
        2: (3240, 6120),
        3: (3024, 1728, 4608),
        4: (3744, 1872, 1872, 1872),
    },
}


# --------------------------------------------------------------------------
# Raw XML helpers.
#
# python-docx has no API for any of this, so these reach into the underlying
# element. Each one is idempotent: it replaces the child it owns rather than
# appending a second copy, so calling them twice is harmless.
# --------------------------------------------------------------------------

def _el(tag: str, **attrs) -> OxmlElement:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), str(v))
    return e


def _sub(parent, tag: str, **attrs) -> OxmlElement:
    e = _el(tag, **attrs)
    parent.append(e)
    return e


def _get_or_add(parent, tag: str) -> OxmlElement:
    found = parent.find(qn(tag))
    if found is None:
        found = OxmlElement(tag)
        parent.append(found)
    return found


def shade(el, fill: str) -> None:
    """Solid background fill on a paragraph's pPr, a cell's tcPr or a run's rPr."""
    old = el.find(qn("w:shd"))
    if old is not None:
        el.remove(old)
    el.append(_el("w:shd", val="clear", color="auto", fill=fill))


def borders(el, tag: str, *, sides: dict) -> None:
    """
    Border set on a paragraph (tag 'w:pBdr'), a table ('w:tblBorders') or a
    cell ('w:tcBorders'). `sides` maps side name -> (size_eighths, colour);
    a side that is absent is simply not written, which is how the callout gets
    a left rule and nothing else.
    """
    old = el.find(qn(tag))
    if old is not None:
        el.remove(old)
    bdr = OxmlElement(tag)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if side in sides:
            sz, color = sides[side]
            _sub(bdr, f"w:{side}", val="single", sz=sz, space=0, color=color)
    el.append(bdr)


def run_shade(run, fill: str) -> None:
    """Character shading - this is what makes the step badge a filled chip."""
    shade(run._r.get_or_add_rPr(), fill)


def cell_margins(tc, top: int, left: int, bottom: int, right: int) -> None:
    tcPr = tc._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcMar"))
    if old is not None:
        tcPr.remove(old)
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        _sub(mar, f"w:{name}", w=val, type="dxa")
    tcPr.append(mar)


def repeat_header(row) -> None:
    """Mark a row as a heading row so Word repeats it on every page."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))


def row_centered(row) -> None:
    """Rows carry their own justification; without it a centred table drifts left."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(_el("w:jc", val="center"))


def fixed_layout(table, widths) -> None:
    """
    Pin the column grid. Word only honours explicit widths under a fixed
    layout, otherwise it re-flows columns to fit the content.
    """
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn("w:tblLayout"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(_el("w:tblLayout", type="fixed"))

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        _sub(grid, "w:gridCol", w=w)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Twips(w)


def field(paragraph, code: str, *, bold=True, color=None, sz=None):
    """
    A Word field such as PAGE or NUMPAGES, written as the begin/instr/separate/
    result/end run sequence. The cached result is a literal '1'; Word and the
    PDF exporter both recompute it while rendering, which is why the build
    never needs to call Fields.Update() (with NUMPAGES that call never returns).
    """
    def _run(children):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if color:
            rPr.append(_el("w:color", val=color))
        if sz:
            rPr.append(_el("w:sz", val=sz))
        r.append(rPr)
        for c in children:
            r.append(c)
        paragraph._p.append(r)

    _run([_el("w:fldChar", fldCharType="begin")])
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    _run([instr])
    _run([_el("w:fldChar", fldCharType="separate")])
    t = OxmlElement("w:t")
    t.text = "1"
    _run([t])
    _run([_el("w:fldChar", fldCharType="end")])


# --------------------------------------------------------------------------
# Run and paragraph helpers.
# --------------------------------------------------------------------------

def add_run(paragraph, text: str, *, bold=None, italic=None, underline=None,
            color=None, sz=None, caps=False, shd=None):
    """
    One formatted run.

    bold/italic/underline default to None and must stay None when the effect is
    not wanted. Passing False writes <w:b w:val="0"/>, which does not mean
    "unstyled" - it actively cancels the bold that the Heading 1/2/3 style is
    supplying, and silently un-bolds every heading in the document.
    """
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = SPEC["font"]
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if sz:
        run.font.size = Pt(sz / 2)
    if caps:
        run._r.get_or_add_rPr().append(OxmlElement("w:caps"))
    if shd:
        run_shade(run, shd)
    return run


# The lookbehind stops a backslash-escaped asterisk from being mistaken for the
# closing delimiter, which matters constantly here: field labels the app marks
# as required end in a literal '*', and they are nearly always bold as well, so
# `**Product Name \***` has to close on the last two asterisks and not the
# first two it happens to run into.
BOLD_RE = re.compile(r"\*\*(.+?)(?<!\\)\*\*", re.S)
ESCAPE_RE = re.compile(r"\\(.)")


def unescape(text: str) -> str:
    """Drop the backslash from an escaped character now that parsing is done."""
    return ESCAPE_RE.sub(r"\1", text)


def add_rich(paragraph, text: str, **base):
    """Render `**bold**` spans as bold runs and everything else with `base`."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            add_run(paragraph, unescape(text[pos:m.start()]), **base)
        bold_kw = dict(base)
        bold_kw["bold"] = True
        add_run(paragraph, unescape(m.group(1)), **bold_kw)
        pos = m.end()
    if pos < len(text):
        add_run(paragraph, unescape(text[pos:]), **base)


def spacing(paragraph, *, before=None, after=None, line=None):
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Twips(before)
    if after is not None:
        pf.space_after = Twips(after)
    if line is not None:
        pf.line_spacing = Twips(line)


def indent(paragraph, *, left=None, right=None, hanging=None):
    pf = paragraph.paragraph_format
    if left is not None:
        pf.left_indent = Twips(left)
    if right is not None:
        pf.right_indent = Twips(right)
    if hanging is not None:
        pf.first_line_indent = Twips(-hanging)


# --------------------------------------------------------------------------
# Document constructs.
# --------------------------------------------------------------------------

class ManualBuilder:
    def __init__(self):
        self.doc = Document()
        self._configure_styles()
        self._configure_page(self.doc.sections[0])
        self.manifest = json.loads(MANIFEST.read_text(encoding="utf-8")) if MANIFEST.exists() else {}
        self.figures_embedded = 0
        self.figures_total = 0
        self.book_subtitle = ""

    # -- setup ------------------------------------------------------------

    def _configure_styles(self):
        """
        Pin the styles the constructs lean on, so the output does not depend on
        whatever python-docx's stock template happens to define.
        """
        st = self.doc.styles

        normal = st["Normal"]
        normal.font.name = SPEC["font"]
        normal.font.size = Pt(SPEC["sz_body"] / 2)
        normal.paragraph_format.space_after = Twips(0)
        rpr = normal.element.get_or_add_rPr()
        rfonts = _get_or_add(rpr, "w:rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rfonts.set(qn(f"w:{attr}"), SPEC["font"])

        # Heading 1/2/3 supply the bold and the outline levels. The outline
        # levels are what Word turns into PDF bookmarks, so the Part/Step tree
        # in the exported PDF comes straight from these three styles.
        for name, sz, lvl in (("Heading 1", 28, 0), ("Heading 2", 26, 1), ("Heading 3", 24, 2)):
            s = st[name]
            s.font.name = SPEC["font"]
            s.font.bold = True
            s.font.size = Pt(sz / 2)
            s.font.color.rgb = RGBColor.from_string(SPEC["accent"])
            ppr = s.element.get_or_add_pPr()
            _get_or_add(ppr, "w:keepNext")
            _get_or_add(ppr, "w:keepLines")
            _get_or_add(ppr, "w:outlineLvl").set(qn("w:val"), str(lvl))

    def _configure_page(self, section):
        section.page_width = Twips(SPEC["page_w"])
        section.page_height = Twips(SPEC["page_h"])
        for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, side, Twips(SPEC["margin"]))
        section.header_distance = Twips(SPEC["hf_distance"])
        section.footer_distance = Twips(SPEC["hf_distance"])

    # -- sections, headers, footers ---------------------------------------

    def new_section(self, *, with_chrome: bool):
        """
        Start a fresh page-break section. `with_chrome` decides whether it
        carries the running header and footer.

        Every section after the first must explicitly unlink its header and
        footer even when it wants them blank -- a linked header silently
        inherits the previous section's, which is exactly how a cover page ends
        up wearing the body chrome.
        """
        section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._configure_page(section)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for part in (section.header, section.footer):
            for p in list(part.paragraphs)[1:]:
                p._element.getparent().remove(p._element)
        if with_chrome:
            self._build_header(section)
            self._build_footer(section)
        else:
            section.header.paragraphs[0].text = ""
            section.footer.paragraphs[0].text = ""
        return section

    def _build_header(self, section):
        p = section.header.paragraphs[0]
        p.text = ""
        borders(p._p.get_or_add_pPr(), "w:pBdr", sides={"bottom": (12, SPEC["rule"])})
        spacing(p, after=80)
        small = dict(sz=SPEC["sz_small"])
        add_run(p, PRODUCT, bold=True, underline=True, color=SPEC["accent"], **small)
        add_run(p, "  |  ", color=SPEC["grey"], **small)
        add_run(p, self.book_subtitle, color=SPEC["grey_dark"], **small)
        p.add_run().add_tab()
        add_run(p, BRAND, bold=True, color=SPEC["accent"], **small)

    def _build_footer(self, section):
        small = dict(sz=SPEC["sz_small"])

        p1 = section.footer.paragraphs[0]
        p1.text = ""
        borders(p1._p.get_or_add_pPr(), "w:pBdr", sides={"top": (12, SPEC["rule"])})
        spacing(p1, before=80)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p1, "Page ", color=SPEC["grey"], **small)
        field(p1, "PAGE", color=SPEC["grey_page"], sz=SPEC["sz_small"])
        add_run(p1, " of ", color=SPEC["grey"], **small)
        field(p1, "NUMPAGES", color=SPEC["grey_page"], sz=SPEC["sz_small"])

        p2 = section.footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, BRAND, bold=True, color=SPEC["accent"], **small)
        add_run(p2, "  ·  ", color=SPEC["grey"], **small)
        add_run(p2, f"{PRODUCT} {self.book_subtitle}", color=SPEC["grey"], **small)

    # -- cover ------------------------------------------------------------

    def cover(self, subtitle: str, tagline: str):
        d = self.doc

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=1200, after=200)
        add_run(p, " ".join(BRAND), bold=True, color=SPEC["accent"], sz=SPEC["sz_brand"])

        self._rule(before=120, after=360)

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=200, after=360)
        add_run(p, PRODUCT, bold=True, caps=True, color=SPEC["navy"], sz=SPEC["sz_title"])

        self._rule(before=160, after=360)

        # Deliberately a plain paragraph rather than the Subtitle style: that
        # style supplies italic, and cancelling it would mean writing
        # <w:i w:val="0"/>, the construct this document never emits.
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=200, after=200)
        add_run(p, subtitle, color=SPEC["accent"], sz=SPEC["sz_subtitle"])

        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=360, after=200)
        add_run(p, tagline, italic=True, color=SPEC["grey"], sz=SPEC["sz_tagline"])

        d.add_paragraph()
        d.add_paragraph()

    def _rule(self, *, before, after):
        p = self.doc.add_paragraph()
        borders(p._p.get_or_add_pPr(), "w:pBdr", sides={"bottom": (12, SPEC["rule"])})
        spacing(p, before=before, after=after)
        return p

    # -- headings ---------------------------------------------------------

    def part_banner(self, text: str):
        """A reversed-out full-bleed bar. The negative indents pull the fill a
        few twips past the text edge on both sides so it reads as a band."""
        p = self.doc.add_paragraph(style="Heading 1")
        pPr = p._p.get_or_add_pPr()
        edge = (2, SPEC["banner_fill"])
        borders(pPr, "w:pBdr", sides={"top": edge, "left": edge, "bottom": edge, "right": edge})
        shade(pPr, SPEC["banner_fill"])
        spacing(p, before=200, after=240)
        indent(p, left=SPEC["banner_indent"], right=SPEC["banner_indent"])
        add_run(p, "  " + text, color="FFFFFF", sz=SPEC["sz_banner"])
        return p

    def heading2(self, text: str):
        p = self.doc.add_paragraph(style="Heading 2")
        spacing(p, before=280, after=120)
        add_run(p, text, color=SPEC["accent"], sz=SPEC["sz_heading"])
        return p

    def step(self, badge: str, title: str):
        p = self.doc.add_paragraph(style="Heading 3")
        spacing(p, before=280, after=120)
        if badge:
            add_run(p, f" {badge} ", color="FFFFFF", sz=SPEC["sz_small"], shd=SPEC["accent"])
            # Bold is left alone here rather than switched off: the spacer is
            # two blank characters, so inherited bold is invisible, whereas
            # cancelling it would poison the Heading 3 bold supply.
            add_run(p, "  ", color=SPEC["body"], sz=SPEC["sz_small"])
        add_run(p, title, color=SPEC["accent"], sz=SPEC["sz_heading"])
        return p

    # -- body -------------------------------------------------------------

    def body(self, text: str):
        p = self.doc.add_paragraph()
        spacing(p, after=160, line=240)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_rich(p, text, color=SPEC["body"])
        return p

    def bullet(self, text: str):
        """
        Hanging-indent bullet with a literal glyph, not a numbering definition.
        A numPr would drag in numbering.xml and Word's list-continuation
        behaviour; this stays inert and always renders identically.
        """
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        _sub(tabs, "w:tab", val="left", pos=SPEC["bullet_indent"])
        pPr.append(tabs)
        spacing(p, after=100)
        indent(p, left=SPEC["bullet_indent"], hanging=SPEC["bullet_hanging"])
        add_run(p, SPEC["bullet_glyph"], bold=True, color=SPEC["accent"])
        p.add_run().add_tab()
        add_rich(p, text, color=SPEC["body"])
        return p

    # -- boxes ------------------------------------------------------------

    def _full_width_box(self, *, border_color: str, fill: str):
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fixed_layout(table, [SPEC["content_w"]])
        row = table.rows[0]
        row_centered(row)
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        borders(tcPr, "w:tcBorders",
                sides={"left": (SPEC["callout_border_sz"], border_color)})
        shade(tcPr, fill)
        cell_margins(cell, 140, 200, 140, 200)
        return table, cell

    def callout(self, label: str, text: str):
        _, cell = self._full_width_box(border_color=SPEC["callout_border"],
                                       fill=SPEC["callout_fill"])
        p = cell.paragraphs[0]
        add_run(p, label, bold=True, italic=True, color=SPEC["accent"])
        add_rich(p, "  " + text, italic=True, color=SPEC["accent"])
        self.doc.add_paragraph()

    def closing_banner(self, text: str):
        _, cell = self._full_width_box(border_color=SPEC["accent"],
                                       fill=SPEC["tbl_head_fill"])
        p = cell.paragraphs[0]
        add_rich(p, text, bold=True, italic=True, color=SPEC["navy"])
        self.doc.add_paragraph()

    def closing_line(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style = dict(italic=True, color=SPEC["grey"], sz=SPEC["sz_closing"])
        add_run(p, "— End of Guide —", **style)
        run = add_run(p, "", **style)
        run._r.append(OxmlElement("w:br"))
        add_run(p, f"{BRAND}  |  {PRODUCT} {self.book_subtitle}", **style)

    # -- tables -----------------------------------------------------------

    def data_table(self, header: list[str], rows: list[list[str]]):
        cols = len(header)
        widths = SPEC["grids"].get(cols)
        if widths is None:
            each = SPEC["content_w"] // cols
            widths = [each] * (cols - 1) + [SPEC["content_w"] - each * (cols - 1)]

        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        edge = (4, SPEC["tbl_border"])
        borders(table._tbl.tblPr, "w:tblBorders",
                sides={k: edge for k in ("top", "left", "bottom", "right", "insideH", "insideV")})
        fixed_layout(table, widths)

        head = table.rows[0]
        row_centered(head)
        repeat_header(head)
        for cell, text in zip(head.cells, header):
            shade(cell._tc.get_or_add_tcPr(), SPEC["tbl_head_fill"])
            cell_margins(cell, 80, 110, 80, 110)
            add_rich(cell.paragraphs[0], text, bold=True, color=SPEC["navy"], sz=SPEC["sz_small"])

        for r, values in enumerate(rows, start=1):
            tr = table.rows[r]
            row_centered(tr)
            for c, (cell, text) in enumerate(zip(tr.cells, values)):
                cell_margins(cell, 80, 110, 80, 110)
                kw = dict(color=SPEC["tbl_col1"] if c == 0 else SPEC["body"], sz=SPEC["sz_small"])
                if c == 0:
                    kw["bold"] = True
                add_rich(cell.paragraphs[0], text, **kw)

        self.doc.add_paragraph()
        return table

    # -- figures ----------------------------------------------------------

    def figure(self, number: str, caption: str):
        """
        A bordered plate holding the screenshot and its caption. If the shot
        named in the manifest is not on disk yet the plate is still drawn, with
        a note in place of the image, so the manual builds complete today and
        improves as captures arrive without a single edit to the Markdown.
        """
        self.figures_total += 1
        entry = self.manifest.get(str(number), {})
        path = SHOTS / entry.get("file", "") if entry.get("file") else None

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        fixed_layout(table, [SPEC["fig_cell_w"]])
        row = table.rows[0]
        row_centered(row)
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        edge = (6, SPEC["fig_border"])
        borders(tcPr, "w:tcBorders",
                sides={"top": edge, "left": edge, "bottom": edge, "right": edge})
        shade(tcPr, SPEC["fig_fill"])
        cell_margins(cell, 260, 180, 260, 180)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=80)

        if path and path.exists():
            w, h = self._fit(path)
            p.add_run().add_picture(str(path), width=w, height=h)
            self.figures_embedded += 1
        else:
            add_run(p, f"[ IMAGE {number} — not captured yet ]",
                    italic=True, color=SPEC["grey"], sz=SPEC["sz_small"])

        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cap, caption, italic=True, color=SPEC["grey"], sz=SPEC["sz_small"])

        self.doc.add_paragraph()
        return table

    @staticmethod
    def _fit(path: Path):
        """Scale to fit the plate, preserving aspect. Portrait screenshots end
        up height-bound and landscape ones width-bound, so both sit on the page
        without one of them blowing the plate open."""
        with Image.open(path) as im:
            iw, ih = im.size
        scale = min(SPEC["fig_max_w_in"] / iw, SPEC["fig_max_h_in"] / ih)
        return Emu(int(Inches(iw * scale))), Emu(int(Inches(ih * scale)))


# --------------------------------------------------------------------------
# Markdown reader.
#
# A deliberately tiny dialect. Anything richer would push formatting decisions
# back into the content file, which is the one thing this split exists to stop.
# --------------------------------------------------------------------------

TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")
STEP_RE = re.compile(r"^(Step\s+\d+)\s+(.*)$")


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def lint_asterisks(text: str) -> list[tuple[int, str]]:
    """
    Find asterisks that will print literally.

    `**bold**` is the only inline mark the dialect knows, so a lone `*` meant as
    italic silently reaches the page as an asterisk. Required-field labels do
    legitimately end in one, but those are written `\\*` -- so once bold spans
    and escapes are removed, anything left over is a mistake worth reporting.
    """
    body = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    out = []
    for n, raw in enumerate(body.splitlines(), 1):
        residue = ESCAPE_RE.sub("", BOLD_RE.sub("", raw))
        if "*" in residue:
            out.append((n, raw.strip()))
    return out


def build():
    if not MD.exists():
        sys.exit(f"No manual source at {MD}")

    b = ManualBuilder()
    source = MD.read_text(encoding="utf-8")
    lines = source.splitlines()

    stray = lint_asterisks(source)
    for n, line in stray:
        print(f"  WARNING line {n}: stray '*' will print literally -- {line[:70]}")

    first_book = True
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # HTML comments span lines, and the block at the top of the source
        # documents this very dialect -- skipping only the opening line would
        # feed its examples in as real content.
        if line.startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if not line:
            i += 1
            continue

        # New book: cover section (no chrome), then a body section that has it.
        if line.startswith("%% BOOK:"):
            subtitle, _, tagline = line[len("%% BOOK:"):].partition("|")
            subtitle, tagline = subtitle.strip(), tagline.strip()
            if not first_book:
                b.new_section(with_chrome=False)
            first_book = False
            b.book_subtitle = subtitle
            b.cover(subtitle, tagline)
            b.new_section(with_chrome=True)
            i += 1
            continue

        if line == "@@ END":
            b.closing_line()
            i += 1
            continue

        if line.startswith("=== "):
            b.closing_banner(line[4:].strip())
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            m = STEP_RE.match(text)
            b.step(m.group(1), m.group(2)) if m else b.step("", text)
            i += 1
            continue

        if line.startswith("## "):
            b.heading2(line[3:].strip())
            i += 1
            continue

        if line.startswith("# "):
            b.part_banner(line[2:].strip())
            i += 1
            continue

        if line.startswith("> "):
            rest = line[2:].strip()
            label, _, text = rest.partition("  ")
            b.callout(label.strip(), text.strip())
            i += 1
            continue

        if line.startswith("- "):
            b.bullet(line[2:].strip())
            i += 1
            continue

        if line.startswith("|"):
            cells = split_row(line)
            # A figure is a one-off two-cell row whose first cell is IMAGE n.
            if len(cells) == 2 and cells[0].upper().startswith("IMAGE "):
                b.figure(cells[0].split()[1], cells[1])
                i += 1
                continue
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                current = lines[i].strip()
                # A figure that follows a table with no blank line between them
                # would otherwise be swallowed as another row.
                nxt = split_row(current)
                if len(nxt) == 2 and nxt[0].upper().startswith("IMAGE "):
                    break
                if not TABLE_SEP.match(current):
                    block.append(nxt)
                i += 1
            if block:
                b.data_table(block[0], block[1:])
            continue

        b.body(line)
        i += 1

    b.doc.save(DOCX)
    print(f"Wrote {DOCX.name}")
    print(f"  sections : {len(b.doc.sections)}")
    print(f"  figures  : {b.figures_embedded} of {b.figures_total} illustrated")


if __name__ == "__main__":
    build()
